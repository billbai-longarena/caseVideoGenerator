from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from server.app.core.config import Settings
from server.app.core.errors import AppError
from server.app.services.contracts import canonical_json
from server.app.services.storage import JobStorage, atomic_write_json, sha256_file, sha256_text
from server.app.services.uploads import UploadStorage


class SourceIngestion:
    """Materialize, extract and normalize source-mode inputs.

    Full source bytes and extracted text stay below the job's local ``source``
    directory.  The model-facing boundary contains only bounded excerpts and
    source references, so downstream callers cannot accidentally send entire
    documents to an external model.
    """

    def __init__(self, settings: Settings, storage: JobStorage, uploads: UploadStorage) -> None:
        self.settings = settings
        self.storage = storage
        self.uploads = uploads

    def ingest(self, job_id: str) -> dict[str, Any]:
        manifest = self.storage.read_manifest(job_id)
        if manifest.get("input_mode") not in {"source", "structured"}:
            raise AppError("source_invalid", "source ingestion requires source or structured input mode")

        source_root = self.storage.job_root(job_id) / "source"
        originals = source_root / "originals"
        extracted = source_root / "extracted"
        originals.mkdir(parents=True, exist_ok=True)
        extracted.mkdir(parents=True, exist_ok=True)

        upload_ids = list(manifest.get("inputs", {}).get("upload_ids", []))
        if len(upload_ids) > self.settings.max_upload_files:
            raise AppError("source_invalid", "too many source files")

        records: list[dict[str, Any]] = []
        extracted_by_source: dict[str, str] = {}
        structured_fields: dict[str, Any] = {}
        total_bytes = 0

        for index, upload_id in enumerate(upload_ids, start=1):
            upload = self.uploads.bind(upload_id, job_id)
            self.uploads.verify_bytes(upload_id)
            if upload["suffix"] == ".zip":
                raise AppError("source_invalid", "zip uploads are only accepted in project mode")
            total_bytes += int(upload["size_bytes"])
            if total_bytes > self.settings.max_upload_bytes:
                raise AppError("source_invalid", "combined source uploads exceed the configured limit", status_code=413)

            source_id = f"src_{index:04d}"
            local_name = f"{source_id}{upload['suffix']}"
            original_path = originals / local_name
            if not original_path.exists():
                shutil.copyfile(self.uploads.data_path(upload_id), original_path)
            if sha256_file(original_path) != upload["sha256"]:
                raise AppError("artifact_corrupt", f"copied source checksum mismatch: {source_id}")

            record = self._base_record(source_id, upload)
            try:
                text, warnings, structured = self._extract(original_path, upload["suffix"])
                text = self._bounded_text(text)
                if not text.strip():
                    if upload["suffix"] == ".pdf":
                        record.update(extraction_status="ocr_required", error_code="source_ocr_required")
                        records.append(record)
                        self._write_manifest(source_root, records)
                        raise AppError(
                            "source_ocr_required",
                            f"PDF appears to be scanned and requires OCR: {upload['filename']}",
                            stage="source.extract",
                        )
                    raise AppError("source_extract_failed", f"no readable text found in {upload['filename']}")
                text_path = extracted / f"{source_id}.txt"
                text_path.write_text(text, encoding="utf-8")
                record.update(
                    extraction_status="succeeded",
                    extracted_text_sha256=sha256_file(text_path),
                    warnings=warnings,
                )
                extracted_by_source[source_id] = text
                if structured is not None:
                    structured_fields[source_id] = structured
            except AppError:
                if record["extraction_status"] == "pending":
                    record.update(extraction_status="failed", error_code="source_extract_failed")
                if record not in records:
                    records.append(record)
                self._write_manifest(source_root, records)
                raise
            except Exception as exc:
                record.update(extraction_status="failed", error_code="source_extract_failed")
                records.append(record)
                self._write_manifest(source_root, records)
                raise AppError(
                    "source_extract_failed",
                    f"failed to extract source: {upload['filename']}",
                    stage="source.extract",
                ) from exc
            records.append(record)

        inline_path = source_root / "structured_input.json"
        if inline_path.is_file():
            payload = json.loads(inline_path.read_text(encoding="utf-8"))
            source_id = f"src_{len(records) + 1:04d}"
            encoded = canonical_json(payload)
            text = self._bounded_text(encoded)
            text_path = extracted / f"{source_id}.txt"
            text_path.write_text(text, encoding="utf-8")
            records.append(
                {
                    "source_id": source_id,
                    "upload_id": None,
                    "original_name": "structured_input.json",
                    "safe_name": "structured_input.json",
                    "media_type": "application/json",
                    "size_bytes": inline_path.stat().st_size,
                    "sha256": sha256_file(inline_path),
                    "extraction_status": "succeeded",
                    "extracted_text_sha256": sha256_file(text_path),
                    "external_sharing_policy": "structured_excerpt",
                    "warnings": [],
                    "error_code": None,
                }
            )
            extracted_by_source[source_id] = text
            structured_fields["inline"] = payload

        if not records:
            raise AppError("source_invalid", "source job has no materialized inputs")

        source_manifest = self._write_manifest(source_root, records)
        case_inputs = self._build_case_inputs(manifest, records, structured_fields)
        self.storage.contracts.validate("case_inputs", "v1", case_inputs, error_code="source_invalid")
        atomic_write_json(source_root / "case_inputs.json", case_inputs)
        boundary = self._build_external_boundary(records, extracted_by_source, structured_fields)
        atomic_write_json(source_root / "external_boundary.json", boundary)
        self.storage.append_event(
            job_id,
            "source.extracted",
            "source.extract",
            "源材料已完成本地提取与外发边界构建",
            {"source_count": len(records), "external_chars": boundary["total_excerpt_chars"]},
        )
        return {"source_manifest": source_manifest, "case_inputs": case_inputs, "boundary": boundary}

    def _base_record(self, source_id: str, upload: dict[str, Any]) -> dict[str, Any]:
        return {
            "source_id": source_id,
            "upload_id": upload["upload_id"],
            "original_name": upload["filename"],
            "safe_name": upload["safe_name"],
            "media_type": upload["detected_media_type"],
            "size_bytes": upload["size_bytes"],
            "sha256": upload["sha256"],
            "extraction_status": "pending",
            "extracted_text_sha256": None,
            "external_sharing_policy": "structured_excerpt",
            "warnings": [],
            "error_code": None,
        }

    def _extract(self, path: Path, suffix: str) -> tuple[str, list[str], Any | None]:
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8-sig"), [], None
        if suffix == ".json":
            payload = json.loads(path.read_text(encoding="utf-8-sig"))
            return canonical_json(payload), [], payload
        if suffix == ".docx":
            document = Document(path)
            blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        blocks.append("\t".join(cells))
            return "\n\n".join(blocks), [], None
        if suffix == ".pdf":
            reader = PdfReader(path, strict=True)
            pages: list[str] = []
            warnings: list[str] = []
            for page_number, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if not text:
                    warnings.append(f"page {page_number} contains no extractable text")
                else:
                    pages.append(text)
                if sum(len(item) for item in pages) > self.settings.max_source_text_chars:
                    raise AppError("source_invalid", "extracted source text exceeds the configured limit")
            return "\n\n".join(pages), warnings, None
        raise AppError("source_invalid", f"unsupported source type: {suffix}")

    def _bounded_text(self, text: str) -> str:
        normalized = text.replace("\x00", "").strip()
        if len(normalized) > self.settings.max_source_text_chars:
            raise AppError("source_invalid", "extracted source text exceeds the configured limit")
        return normalized

    def _write_manifest(self, source_root: Path, records: list[dict[str, Any]]) -> dict[str, Any]:
        payload = {"version": "1", "files": records}
        self.storage.contracts.validate("source_manifest", "v1", payload, error_code="source_invalid")
        atomic_write_json(source_root / "source_manifest.json", payload)
        return payload

    def _build_case_inputs(
        self,
        manifest: dict[str, Any],
        records: list[dict[str, Any]],
        structured_fields: dict[str, Any],
    ) -> dict[str, Any]:
        duration = manifest.get("target_duration")
        seconds = manifest.get("target_duration_seconds")
        if seconds:
            duration = f"{seconds['min']}-{seconds['max']} seconds"
        return {
            "version": "1",
            "project_name": manifest["project_name"],
            "source_refs": [record["source_id"] for record in records],
            "production": {
                "column": manifest.get("program", "销售不复杂"),
                "target_duration": duration,
                "approval_mode": manifest["approval_mode"],
            },
            "structured_fields": structured_fields,
        }

    def _build_external_boundary(
        self,
        records: list[dict[str, Any]],
        extracted_by_source: dict[str, str],
        structured_fields: dict[str, Any],
    ) -> dict[str, Any]:
        remaining = self.settings.max_external_excerpt_chars
        sources: list[dict[str, Any]] = []
        for record in records:
            text = extracted_by_source.get(record["source_id"], "")
            allowance = min(1_200, remaining)
            excerpt = text[:allowance] if allowance > 0 else ""
            remaining -= len(excerpt)
            sources.append(
                {
                    "source_id": record["source_id"],
                    "policy": record["external_sharing_policy"],
                    "excerpt": excerpt,
                    "excerpt_sha256": sha256_text(excerpt),
                    "truncated": len(text) > len(excerpt),
                }
            )
        total = sum(len(item["excerpt"]) for item in sources)
        if total > self.settings.max_external_excerpt_chars:
            raise AppError("contract_invalid", "external source boundary exceeded configured limit")
        return {
            "version": "1",
            "sources": sources,
            # Values are already represented by bounded excerpts above. Only
            # disclose shape metadata here; the original object could bypass
            # the external-sharing ceiling.
            "structured_fields": self._structured_shape(structured_fields),
            "total_excerpt_chars": total,
            "max_excerpt_chars": self.settings.max_external_excerpt_chars,
        }

    def _structured_shape(self, fields: dict[str, Any]) -> dict[str, Any]:
        result: dict[str, Any] = {}
        for source_key, value in fields.items():
            if isinstance(value, dict):
                keys = [str(key)[:80] for key in list(value)[:50]]
                result[source_key] = {"type": "object", "keys": keys, "truncated": len(value) > len(keys)}
            elif isinstance(value, list):
                result[source_key] = {"type": "array", "item_count": len(value)}
            else:
                result[source_key] = {"type": type(value).__name__}
        return result
